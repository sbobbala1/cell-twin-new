from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "output" / "docx" / "human_cell_digital_twin_process_library.docx"


PROCESSES = [
    ("Process 1", "Membrane Transport", "p01_transport.py"),
    ("Process 2", "Glycolysis", "p02_glycolysis.py"),
    ("Process 3", "Oxidative Phosphorylation", "p03_oxphos.py"),
    ("Process 4", "Fermentation", "p04_fermentation.py"),
    ("Process 5", "ATP Demand", "p05_atp_demand.py"),
    ("Process 6", "HIF-1a Oxygen Sensing", "p06_hif1a.py"),
    ("Process 7", "Calcium Signalling", "p07_calcium.py"),
    ("Process 8", "mTOR Protein Synthesis", "p08_mtor.py"),
    ("Process 9", "Na/K-ATPase Pump", "p09_nakpump.py"),
]


SUMMARIES = {
    "Process 1": {
        "description": "Models membrane-level glucose exchange, oxygen entry, and passive sodium leak. GLUT is now reversible, so starvation can drain cytosolic glucose when external glucose is zero.",
        "formulas": [
            "v_glut = Vmax_glut * glut_scale * (mm(Glc_ext) - mm(Glc_cyt))",
            "v_o2 = Vmax_o2 * O2_ext / (Km_o2 + O2_ext)",
            "glut_scale = 1 + (HIF_MAX_GLUT - 1) * hif1a",
        ],
    },
    "Process 2": {
        "description": "Converts glucose to pyruvate and net ATP. The reaction is glucose-limited and ADP-limited.",
        "formulas": [
            "v_glyc = Vmax_glyc * mm(Glucose, Km_glc) * mm(ADP, Km_adp)",
            "Glucose + 2 ADP -> 2 Pyruvate + 2 ATP",
        ],
    },
    "Process 3": {
        "description": "Uses pyruvate and oxygen to generate the large ATP payoff of aerobic metabolism. The reaction is pyruvate-, oxygen-, and ADP-limited.",
        "formulas": [
            "v_oxphos = Vmax_oxphos * mm(Pyruvate, Km_pyr) * mm(O2, Km_o2)",
            "Pyruvate + 3 O2 + 30 ADP -> 30 ATP",
        ],
    },
    "Process 4": {
        "description": "Converts pyruvate to lactate under low oxygen. HIF-1a raises fermentation capacity during hypoxia.",
        "formulas": [
            "o2_inhibition = 1 - O2 / (Km_o2_inh + O2)",
            "v_ferm = Vmax_ferm * ferm_scale * mm(Pyruvate, Km_pyr) * o2_inhibition",
            "ferm_scale = 1 + (HIF_MAX_FERM - 1) * hif1a",
        ],
    },
    "Process 5": {
        "description": "Represents baseline housekeeping ATP usage from many processes not modeled explicitly.",
        "formulas": [
            "v_demand = ATP_DEMAND_BASE * ATP / (Km_demand + ATP)",
            "ATP -> ADP",
        ],
    },
    "Process 6": {
        "description": "Models the hypoxia switch. When oxygen falls below the threshold, HIF-1a rises and gates glucose import and fermentation.",
        "formulas": [
            "if O2 < HIF_O2_THRESH: dHIF/dt = HIF_RISE_RATE * (1 - HIF)",
            "else: dHIF/dt = -HIF_FALL_RATE * HIF",
            "HIF is clamped to [0, 1]",
        ],
    },
    "Process 7": {
        "description": "Keeps cytosolic calcium near its resting level with a simple leak-plus-pump homeostasis model.",
        "formulas": [
            "v_leak = CA_LEAK_RATE",
            "v_pump = CA_PUMP_RATE * max(Ca_cyt - CA_REST, 0)",
            "dCa/dt = v_leak - v_pump",
        ],
    },
    "Process 8": {
        "description": "Links ATP and amino acid availability to protein synthesis through an mTOR-like anabolic gate.",
        "formulas": [
            "mTOR_activity = clamp(ATP / NORMAL_ATP, 0, 1)",
            "v_synth = Vmax_ribo * mTOR_activity * mm(ATP, Km_atp) * mm(AA, Km_aa)",
            "4 ATP + 1 amino acid -> protein_mass + 4 ADP",
        ],
    },
    "Process 9": {
        "description": "Models Na/K-ATPase as a major ATP consumer that maintains sodium and potassium gradients.",
        "formulas": [
            "v_pump = Vmax_pump * pump_activity * mm(ATP) * Hill(Na_cyt, Km_Na, 3) * Hill(K_ext, Km_K, 2)",
            "1 ATP + 3 Na_cyt + 2 K_ext -> ADP + 3 Na_ext + 2 K_cyt",
        ],
    },
}


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in [("top", top), ("start", start), ("bottom", bottom), ("end", end)]:
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_table_width(table, widths):
    table.autofit = False
    for row in table.rows:
        for idx, width in enumerate(widths):
            cell = row.cells[idx]
            cell.width = width
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(int(width.inches * 1440)))
            tc_w.set(qn("w:type"), "dxa")


def configure_styles(doc):
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    for name, size, color, before, after in [
        ("Heading 1", 16, "2E74B5", 18, 10),
        ("Heading 2", 13, "2E74B5", 14, 7),
        ("Heading 3", 12, "1F4D78", 10, 5),
    ]:
        style = styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    title = styles["Title"]
    title.font.name = "Calibri"
    title.font.size = Pt(24)
    title.font.bold = True
    title.font.color.rgb = RGBColor.from_string("0B2545")
    title.paragraph_format.space_after = Pt(3)

    subtitle = styles.add_style("DocSubtitle", WD_STYLE_TYPE.PARAGRAPH)
    subtitle.font.name = "Calibri"
    subtitle.font.size = Pt(11)
    subtitle.font.color.rgb = RGBColor.from_string("555555")
    subtitle.paragraph_format.space_after = Pt(12)

    code = styles.add_style("CodeBlock", WD_STYLE_TYPE.PARAGRAPH)
    code.font.name = "Courier New"
    code.font.size = Pt(8)
    code.paragraph_format.space_before = Pt(3)
    code.paragraph_format.space_after = Pt(8)
    code.paragraph_format.line_spacing = 1.0
    code.paragraph_format.left_indent = Inches(0.12)

    footer = section.footer.paragraphs[0]
    footer.text = "Human Cell Digital Twin - Process Library"
    footer.style = normal
    footer.runs[0].font.size = Pt(8)
    footer.runs[0].font.color.rgb = RGBColor.from_string("666666")


def add_bullet(doc, text):
    p = doc.add_paragraph(text, style="List Bullet")
    p.paragraph_format.left_indent = Inches(0.375)
    p.paragraph_format.first_line_indent = Inches(-0.188)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.25
    return p


def add_number(doc, text):
    p = doc.add_paragraph(text, style="List Number")
    p.paragraph_format.left_indent = Inches(0.375)
    p.paragraph_format.first_line_indent = Inches(-0.188)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.25
    return p


def add_code(doc, text):
    p = doc.add_paragraph(style="CodeBlock")
    p.paragraph_format.keep_together = True
    run = p.add_run(text)
    run.font.name = "Courier New"
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor.from_string("111111")
    p._p.get_or_add_pPr().append(_paragraph_shading("F4F6F9"))


def _paragraph_shading(fill):
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    return shd


def source_excerpt(filename):
    text = (ROOT / filename).read_text(encoding="utf-8")
    lines = []
    in_compute = False
    for line in text.splitlines():
        if line.strip().startswith("def compute"):
            in_compute = True
        if in_compute:
            lines.append(line)
        if in_compute and line.strip() == "return s":
            break
    return "\n".join(lines).strip()


def add_status_table(doc):
    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for idx, text in enumerate(["Area", "Current state", "Next action"]):
        hdr[idx].text = text
        set_cell_shading(hdr[idx], "E8EEF5")
    rows = [
        ("ODE model", "Runnable; validate.py reports 10/10.", "Keep this as the stable baseline."),
        ("Hypoxia demo", "O2 cut raises HIF-1a, lactate rises, ATP recovers after O2 returns.", "Use this as the first visual scenario."),
        ("Starvation", "Fixed by reversible GLUT and ADP-limited ATP production.", "Refine with better nutrient pools later."),
        ("Spatial model", "Not implemented yet.", "Start with 1-D O2 diffusion, then 32^3 NumPy."),
        ("Visualization", "plot.py exists, but matplotlib is not installed in base Python.", "Install deps or use bundled runtime for simulation."),
    ]
    for row_data in rows:
        row = table.add_row().cells
        for idx, text in enumerate(row_data):
            row[idx].text = text
    set_table_width(table, [Inches(1.35), Inches(2.85), Inches(2.3)])
    for row in table.rows:
        for cell in row.cells:
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)
            for p in cell.paragraphs:
                p.paragraph_format.space_after = Pt(0)
                for run in p.runs:
                    run.font.size = Pt(9)


def build_docx():
    doc = Document()
    configure_styles(doc)

    title = doc.add_paragraph("Human Cell Digital Twin", style="Title")
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    doc.add_paragraph("Core process library, formulas, code blocks, and continuation plan", style="DocSubtitle")

    doc.add_heading("Project Summary", level=1)
    doc.add_paragraph(
        "You built a minimum viable biology/math layer for a human-cell digital twin. The current code is a compartmental ODE simulation: one state dictionary stores concentrations, each process computes a flux for one time step, and simulate.py applies the nine process modules in biological order."
    )
    doc.add_paragraph(
        "The main conceptual upgrade is coupling oxygen stress to downstream behavior. Low O2 stabilizes HIF-1a, and HIF-1a increases GLUT-mediated glucose exchange and fermentation capacity. The code now also passes all validation checks after fixing starvation behavior."
    )

    doc.add_heading("Current Status", level=1)
    add_status_table(doc)

    doc.add_heading("How The Code Is Organized", level=1)
    for item in [
        "constants.py centralizes kinetic rates, Km values, time step, and simulation duration.",
        "state.py defines the initial cytosolic, extracellular, ion, oxygen, energy, amino acid, protein, HIF-1a, and calcium pools.",
        "helpers.py provides clamp, Michaelis-Menten saturation, and Hill kinetics.",
        "p01 through p09 are the process modules. Each has compute(...) and apply(...) functions.",
        "simulate.py controls scenarios, process order, time stepping, and history recording.",
        "validate.py runs biological sanity checks for normal, hypoxia, and starvation scenarios.",
    ]:
        add_bullet(doc, item)

    doc.add_heading("Core Process Library", level=1)
    for proc_id, title_text, filename in PROCESSES:
        info = SUMMARIES[proc_id]
        doc.add_heading(f"{proc_id}: {title_text}", level=2)
        doc.add_paragraph(info["description"])
        doc.add_paragraph("Core formulas:")
        for formula in info["formulas"]:
            add_bullet(doc, formula)
        doc.add_paragraph("Code excerpt:")
        add_code(doc, source_excerpt(filename))

    doc.add_section(WD_SECTION_START.NEW_PAGE)
    doc.add_heading("Spatial Extension Plan", level=1)
    doc.add_paragraph(
        "The next architecture is reaction-diffusion. Instead of one scalar concentration per molecule, each molecule becomes a 3-D array over voxels. Each voxel runs the same local reaction equations and exchanges material with its six orthogonal neighbors through a finite-difference Laplacian."
    )
    for formula in [
        "dC/dt = D * Laplacian(C) + R(C)",
        "Laplacian(C) = (C[x+1] + C[x-1] + C[y+1] + C[y-1] + C[z+1] + C[z-1] - 6*C[x,y,z]) / dx^2",
        "dATP/dt = D_ATP*Lap(ATP) + v_oxphos - v_demand - v_synth - v_pump",
        "dO2/dt = D_O2*Lap(O2) + v_import - 3*v_oxphos",
        "dGlucose/dt = D_Glc*Lap(Glucose) + v_glut - v_glycolysis",
        "dLactate/dt = D_Lac*Lap(Lactate) + v_fermentation - v_export",
    ]:
        add_bullet(doc, formula)

    doc.add_heading("Next Steps", level=1)
    for item in [
        "Refine O2 exchange so oxygen is also gradient-driven instead of one-way import.",
        "Add a requirements.txt with numpy, matplotlib, reportlab, and python-docx for local reproducibility.",
        "Install matplotlib in the environment you want to use, then run plot.py.",
        "Build a tiny 1-D O2 diffusion prototype using Fick's second law.",
        "Promote state variables to a 32^3 NumPy grid and vectorize the existing reactions over voxels.",
        "Add membrane and cytosol masks so transport only happens at the cell boundary.",
        "Move the Laplacian stencil to CuPy, CUDA, or WGSL using ping-pong buffers.",
        "Connect the output buffers to a three.js/WebGL volume renderer for ATP, O2, lactate, Na_cyt, and protein_mass.",
    ]:
        add_number(doc, item)

    doc.add_heading("Validation Command", level=1)
    add_code(doc, "/Users/saha/.cache/codex-runtimes/codex-primary-runtime/dependencies/python/bin/python3 validate.py")

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    return OUT


if __name__ == "__main__":
    print(build_docx())
